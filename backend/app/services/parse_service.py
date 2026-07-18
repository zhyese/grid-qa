"""文档解析：数字文档文本提取 + 扫描件/图片 OCR。

OCR 引擎说明：Windows + paddlepaddle 3.3.1 原生 PaddleOCR 存在 onednn PIR
`ConvertPirAttribute2RuntimeAttribute` 未实现的引擎 bug（实测关闭 oneDNN / 关闭 PIR
新执行器 / monkey-patch enable_mkldnn 均无效，属性内嵌于模型）。故改用
rapidocr-onnxruntime —— 基于 PaddleOCR 官方 PP-OCR 模型、onnxruntime 推理后端，
识别效果一致。生产 Linux/Docker 环境可切回原生 paddleocr（无此 bug）。
"""
import io
from typing import Tuple

_OCR_INSTANCE = None


def _get_ocr():
    global _OCR_INSTANCE
    if _OCR_INSTANCE is None:
        from rapidocr_onnxruntime import RapidOCR

        _OCR_INSTANCE = RapidOCR()
    return _OCR_INSTANCE


def _ext(name: str) -> str:
    return "." + name.rsplit(".", 1)[-1].lower() if "." in name else ""


# ---------- 数字文档 ----------


def extract_pdf(content: bytes) -> Tuple[str, bool]:
    """每页平均文本<10字 → 判定扫描件，走 OCR。"""
    import pdfplumber

    parts, n = [], 0
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        n = len(pdf.pages)
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    text = "\n".join(parts).strip()
    is_scanned = n > 0 and len(text) < n * 10
    return text, is_scanned


def extract_docx(content: bytes) -> str:
    from docx import Document

    return "\n".join(p.text for p in Document(io.BytesIO(content)).paragraphs).strip()


def extract_txt(content: bytes) -> str:
    for enc in ("utf-8", "gbk", "gb2312"):
        try:
            return content.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore").strip()


# ---------- 扫描件 / 图片 OCR（rapidocr-onnxruntime）----------


def _extract_ocr_texts(result) -> list:
    """rapidocr 返回 (result, elapse)；result=list[[box,text,score]] 或 None。"""
    texts = []
    if not result:
        return texts
    for item in result:
        if item and len(item) >= 2:
            texts.append(item[1])
    return texts


def _ocr_ndarray(img) -> str:
    result, _elapse = _get_ocr()(img)
    return "\n".join(_extract_ocr_texts(result))


def ocr_image(content: bytes) -> str:
    import numpy as np
    from PIL import Image

    img = np.array(Image.open(io.BytesIO(content)).convert("RGB"))
    return _ocr_ndarray(img)


def ocr_pdf(content: bytes) -> str:
    import fitz
    import numpy as np
    from PIL import Image

    doc = fitz.open(stream=content, filetype="pdf")
    lines = []
    for page in doc:
        pix = page.get_pixmap(dpi=200)
        img = np.array(Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB"))
        lines.append(_ocr_ndarray(img))
    return "\n".join(s for s in lines if s)


# ---------- 路由 ----------


def parse_file(name: str, content: bytes) -> Tuple[str, bool]:
    """按扩展名路由；返回 (text, is_scanned)。is_scanned=True 表示需走 OCR。"""
    ext = _ext(name)
    if ext == ".pdf":
        return extract_pdf(content)
    if ext in (".doc", ".docx"):
        return extract_docx(content), False
    if ext in (".txt", ".md"):
        return extract_txt(content), False
    if ext in (".png", ".jpg", ".jpeg"):
        return "", True
    raise ValueError(f"不支持的文件类型：{ext}")


def ocr_by_name(name: str, content: bytes) -> str:
    return ocr_pdf(content) if _ext(name) == ".pdf" else ocr_image(content)


# ---------- 结构化解析（S4 升级：表格保留为 markdown，Excel/Word 表格，供结构感知分块）----------
# 旧版 extract_* 返回一整坨 text，表格被打碎成乱序文本。新版返回结构化段落列表：
#   [{type: "text"|"table", content}]，表格整体不被切，检索/生成时结构完整。


def _table_to_markdown(rows: list) -> str:
    """二维单元格列表 → markdown 表格。空单元格转空串，过滤全空行，补齐列宽。"""
    clean = []
    for row in rows:
        cells = ["" if c is None else str(c).strip() for c in (row or [])]
        if any(cells):
            clean.append(cells)
    if not clean:
        return ""
    width = max(len(r) for r in clean)
    clean = [r + [""] * (width - len(r)) for r in clean]
    head, body = clean[0], clean[1:]
    lines = [
        "| " + " | ".join(head) + " |",
        "|" + "|".join(["---"] * width) + "|",
    ]
    for r in body:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


def _first_row_as_header(rows: list) -> str:
    """表格首行作表头（table_header 字段，防数值丢列上下文）。

    可核验引用增强：检索命中数值行时，table_header 回带列名上下文，
    让 LLM 能回答「85 度」对应的设备/限值含义，无需额外回表。
    """
    if not rows:
        return ""
    head = ["" if c is None else str(c).strip() for c in (rows[0] or [])]
    return " | ".join(h for h in head if h)


def extract_xlsx(content: bytes) -> list[dict]:
    """Excel 多 sheet → 每 sheet 一张 markdown 表格段落（运维台账/定值单常见载体）。

    可核验引用增强：表格段带 table_header（首行列名），供 chunk 透传到引用展示。
    """
    from openpyxl import load_workbook

    sections: list[dict] = []
    wb = load_workbook(io.BytesIO(content), data_only=True, read_only=True)
    for ws in wb.worksheets:
        rows = list(ws.iter_rows(values_only=True))
        md = _table_to_markdown(rows)
        if md.strip():
            sections.append({"type": "table", "content": f"## {ws.title}\n{md}",
                             "table_header": _first_row_as_header(rows)})
    return sections


def extract_pdf_structured(content: bytes) -> Tuple[list[dict], bool]:
    """PDF 结构化：每页表格转 markdown + 正文文本。表格整体成段，不再被打碎。

    可核验引用增强：每个 section 带 page_num + 首字符 bbox（前端 PDF 高亮锚点）。
    """
    import json
    import pdfplumber

    sections: list[dict] = []
    n_pages, total_chars = 0, 0
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            n_pages += 1
            page_num = page.page_number
            for tbl in page.extract_tables() or []:
                md = _table_to_markdown(tbl)
                if md.strip():
                    sections.append({"type": "table", "content": md,
                                     "page_num": page_num,
                                     "table_header": _first_row_as_header(tbl)})
            txt = (page.extract_text() or "").strip()
            if txt:
                # 首字符 bbox 作高亮锚点（页内首字矩形，前端 PDF 定位用）
                bbox = None
                try:
                    chars = page.chars
                    if chars:
                        x0, top = chars[0]["x0"], chars[0]["top"]
                        # 取首行（前 40 字符）右下角，避免单字框过窄不可见
                        head = chars[:min(len(chars), 40)]
                        x1 = max(c["x1"] for c in head)
                        bottom = max(c["bottom"] for c in head)
                        bbox = json.dumps([round(x0, 1), round(top, 1),
                                           round(x1, 1), round(bottom, 1)])
                except Exception:
                    bbox = None
                sections.append({"type": "text", "content": txt,
                                 "page_num": page_num, "bbox": bbox})
                total_chars += len(txt)
    is_scanned = n_pages > 0 and total_chars < n_pages * 10
    return sections, is_scanned


def extract_docx_structured(content: bytes) -> list[dict]:
    """Word：表格转 markdown + 段落正文（保留规程表格结构）。

    可核验引用增强：表格段带 table_header（首行列名）。
    """
    from docx import Document

    doc = Document(io.BytesIO(content))
    sections: list[dict] = []
    for tbl in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
        md = _table_to_markdown(rows)
        if md.strip():
            sections.append({"type": "table", "content": md,
                             "table_header": _first_row_as_header(rows)})
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if paras:
        sections.append({"type": "text", "content": "\n".join(paras)})
    return sections


def parse_file_structured(name: str, content: bytes) -> Tuple[list[dict], bool]:
    """结构化解析路由。返回 (sections, is_scanned)。

    sections: [{type:"text"|"table", content}]；is_scanned=True 表示需走 OCR。
    旧 parse_file 保留兼容（仍返回整坨 text）。
    """
    ext = _ext(name)
    if ext == ".pdf":
        return extract_pdf_structured(content)
    if ext in (".doc", ".docx"):
        return extract_docx_structured(content), False
    if ext == ".xlsx":
        return extract_xlsx(content), False
    if ext in (".txt", ".md"):
        return [{"type": "text", "content": extract_txt(content)}], False
    if ext in (".png", ".jpg", ".jpeg"):
        return [], True  # 扫描件/图片 → OCR
    raise ValueError(f"不支持的文件类型：{ext}")


def ocr_to_sections(name: str, content: bytes) -> list[dict]:
    """扫描件/图片 OCR 后包成结构化段落（纯文本走 text 段）。"""
    text = ocr_by_name(name, content)
    return [{"type": "text", "content": text}] if text.strip() else []
