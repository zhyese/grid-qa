"""答案导出：问答 → Word(.docx) 运维报告，供现场打印归档。

复用 python-docx（已在依赖）。把 问题/答复/引用来源/可信度/安全提示 落成结构化文档。
"""
import io
from datetime import datetime

_CONF_MAP = {
    "high": "✓ 高（证据充分，答案可信）",
    "medium": "⚠ 中（证据有限，部分内容建议人工核对）",
    "refused": "✗ 低（证据不足，已保守处理）",
}


def build_docx(query: str, answer: str, sources: list, meta: dict | None = None) -> bytes:
    """生成运维问答报告 .docx，返回字节流。"""
    from docx import Document

    doc = Document()
    doc.add_heading("电网运维智能问答报告", level=0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("一、问题", level=1)
    doc.add_paragraph(query or "")

    doc.add_heading("二、答复", level=1)
    for line in (answer or "").split("\n"):
        line = line.strip()
        if line:
            doc.add_paragraph(line)

    doc.add_heading("三、引用来源", level=1)
    if sources:
        for i, s in enumerate(sources, 1):
            text = s.get("text", "") if isinstance(s, dict) else str(s)
            doc_name = s.get("docName", "") if isinstance(s, dict) else ""
            doc.add_paragraph(f"[{i}] {doc_name}：{text}".strip("： "), style="List Bullet")
    else:
        doc.add_paragraph("无")

    doc.add_heading("四、可信度", level=1)
    meta = meta or {}
    conf = meta.get("confidence", "")
    doc.add_paragraph(f"置信度：{_CONF_MAP.get(conf, conf or '未评估')}")
    if meta.get("hallucinationRate") is not None:
        doc.add_paragraph(f"幻觉率：{meta['hallucinationRate']}")
    if meta.get("responseTime") is not None:
        doc.add_paragraph(f"耗时：{meta['responseTime']} 秒")

    doc.add_paragraph()
    doc.add_paragraph("⚠ 安全提示：现场操作前必须核对调度指令与安规，本报告仅供辅助参考。")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_xlsx(query: str, answer: str, sources: list, meta: dict | None = None) -> bytes:
    """生成问答报告 .xlsx（结构化表格，便于台账登记/二次处理）。"""
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "问答报告"
    ws.append(["字段", "内容"])
    ws.append(["问题", query or ""])
    ws.append(["答复", (answer or "")[:32000]])   # Excel 单元格上限
    ws.append(["生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
    m = meta or {}
    ws.append(["置信度", m.get("confidence", "")])
    ws.append(["耗时(秒)", m.get("responseTime", "")])
    ws.append([])
    ws.append(["引用来源"])
    ws.append(["序号", "文档", "内容"])
    for i, s in enumerate(sources or [], 1):
        text = (s.get("text", "") if isinstance(s, dict) else str(s))[:32000]
        name = (s.get("docName", "") if isinstance(s, dict) else "")
        ws.append([i, name, text])
    ws.column_dimensions["A"].width = 12
    ws.column_dimensions["B"].width = 28
    ws.column_dimensions["C"].width = 80
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
